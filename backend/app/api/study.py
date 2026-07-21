"""Highlighting, annotations, and extraction endpoints."""

import json

from fastapi import APIRouter, Depends, HTTPException, Query
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.database import get_session
from app.models.models import Annotation, Document, Page

router = APIRouter()


class AnnotationCreate(BaseModel):
    page_index: int = 0
    type: str = "highlight"
    bbox: str = "{}"
    color: str = "#FFFF00"
    label: str = ""
    note_text: str = ""
    question_type: str = ""
    options: str = "[]"


@router.post("/{document_id}/annotations")
async def create_annotation(
    document_id: str,
    body: AnnotationCreate,
    session: AsyncSession = Depends(get_session),
):
    """Add a highlight, note, or bookmark to a document."""
    # Verify document exists
    result = await session.execute(
        select(Document).where(Document.id == document_id)
    )
    if not result.scalar_one_or_none():
        raise HTTPException(404, "Document not found")

    annotation = Annotation(
        document_id=document_id,
        page_index=body.page_index,
        type=body.type,
        bbox=body.bbox,
        color=body.color,
        label=body.label,
        note_text=body.note_text,
        question_type=body.question_type,
        options=body.options,
    )
    session.add(annotation)
    await session.commit()
    await session.refresh(annotation)

    return {
        "id": annotation.id,
        "document_id": annotation.document_id,
        "page_index": annotation.page_index,
        "type": annotation.type,
        "bbox": annotation.bbox,
        "color": annotation.color,
        "label": annotation.label,
        "note_text": annotation.note_text,
        "question_type": annotation.question_type,
        "options": annotation.options,
        "created_at": annotation.created_at.isoformat() if annotation.created_at else None,
    }


@router.delete("/{document_id}/annotations/{annotation_id}")
async def delete_annotation(
    document_id: str,
    annotation_id: str,
    session: AsyncSession = Depends(get_session),
):
    """Remove an annotation."""
    result = await session.execute(
        select(Annotation).where(
            Annotation.id == annotation_id,
            Annotation.document_id == document_id,
        )
    )
    annotation = result.scalar_one_or_none()
    if not annotation:
        raise HTTPException(404, "Annotation not found")

    await session.delete(annotation)
    await session.commit()

    return {"deleted": True, "annotation_id": annotation_id}


@router.post("/{document_id}/extract")
async def extract_highlights(
    document_id: str,
    colors: str = Query(default="", description="Comma-separated colors to include"),
    session: AsyncSession = Depends(get_session),
):
    """Extract highlights to a study guide format."""
    # Verify document exists
    result = await session.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    # Get annotations
    color_filter = [c.strip() for c in colors.split(",") if c.strip()] if colors else []
    query = select(Annotation).where(Annotation.document_id == document_id)
    if color_filter:
        query = query.where(Annotation.color.in_(color_filter))
    query = query.order_by(Annotation.page_index, Annotation.created_at)

    result = await session.execute(query)
    annotations = result.scalars().all()

    # Get all pages for context
    result = await session.execute(
        select(Page).where(Page.document_id == document_id).order_by(Page.page_index)
    )
    pages = {(p.page_index): p.text for p in result.scalars().all()}

    # Build study guide
    study_guide = []
    for ann in annotations:
        entry = {
            "annotation_id": ann.id,
            "page_index": ann.page_index,
            "type": ann.type,
            "color": ann.color,
            "label": ann.label,
            "note_text": ann.note_text,
            "bbox": ann.bbox,
        }
        if ann.page_index in pages:
            entry["page_text_snippet"] = pages[ann.page_index][:200]
        study_guide.append(entry)

    return {
        "document_id": document_id,
        "document_title": doc.title,
        "extraction_type": "study_guide",
        "annotations_count": len(study_guide),
        "annotations": study_guide,
    }


@router.get("/{document_id}/export")
async def export_document(
    document_id: str,
    format: str = Query(default="txt"),
    session: AsyncSession = Depends(get_session),
):
    """Export document with annotations to TXT, PDF, or DOCX."""
    # Verify document exists
    result = await session.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    # Get pages
    result = await session.execute(
        select(Page)
        .where(Page.document_id == document_id)
        .order_by(Page.page_index)
    )
    pages = result.scalars().all()

    # Get annotations
    result = await session.execute(
        select(Annotation)
        .where(Annotation.document_id == document_id)
        .order_by(Annotation.page_index)
    )
    annotations = result.scalars().all()

    if format == "txt":
        return _export_txt(doc, pages, annotations)
    elif format == "docx":
        return _export_docx(doc, pages, annotations)
    elif format == "pdf":
        return _export_pdf(doc, pages, annotations)
    else:
        raise HTTPException(400, f"Unsupported export format: {format}")


def _export_txt(doc: Document, pages: list[Page], annotations: list[Annotation]) -> dict:
    """Export as plain text."""
    lines = [f"# {doc.title}", f"Author: {doc.author or 'Unknown'}", ""]
    ann_by_page: dict[int, list[Annotation]] = {}
    for a in annotations:
        ann_by_page.setdefault(a.page_index, []).append(a)

    for page in pages:
        lines.append(f"--- Page {page.page_index + 1} ---")
        lines.append(page.text)
        for ann in ann_by_page.get(page.page_index, []):
            if ann.note_text:
                lines.append(f"  [{ann.type}/{ann.label}] {ann.note_text}")
        lines.append("")

    return {"format": "txt", "filename": f"{doc.title}.txt", "content": "\n".join(lines)}


def _export_docx(doc: Document, pages: list[Page], annotations: list[Annotation]) -> dict:
    """Export as DOCX."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor

        d = DocxDocument()
        d.add_heading(doc.title, 0)
        if doc.author:
            d.add_paragraph(f"Author: {doc.author}")

        ann_by_page: dict[int, list[Annotation]] = {}
        for a in annotations:
            ann_by_page.setdefault(a.page_index, []).append(a)

        for page in pages:
            d.add_heading(f"Page {page.page_index + 1}", level=1)
            d.add_paragraph(page.text)
            for ann in ann_by_page.get(page.page_index, []):
                if ann.note_text:
                    p = d.add_paragraph()
                    run = p.add_run(f"[{ann.type}/{ann.label}] {ann.note_text}")
                    run.font.size = Pt(10)
                    run.font.italic = True

        import tempfile, base64, os
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        d.save(tmp_path)
        with open(tmp_path, "rb") as f:
            data = base64.b64encode(f.read()).decode()
        os.unlink(tmp_path)

        return {"format": "docx", "filename": f"{doc.title}.docx", "content_base64": data}
    except ImportError:
        raise HTTPException(501, "python-docx not available for DOCX export")


def _export_pdf(doc: Document, pages: list[Page], annotations: list[Annotation]) -> dict:
    """Export as PDF (simple text-based via fpdf2 or fallback)."""
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)

        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, doc.title, ln=True)
        pdf.ln(5)

        ann_by_page: dict[int, list[Annotation]] = {}
        for a in annotations:
            ann_by_page.setdefault(a.page_index, []).append(a)

        pdf.set_font("Helvetica", size=11)
        for page in pages:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, f"Page {page.page_index + 1}", ln=True)
            pdf.set_font("Helvetica", size=11)
            # Split long text into multi_cell
            for line in page.text.split("\n"):
                pdf.multi_cell(0, 5, line[:200])  # Truncate for PDF
            for ann in ann_by_page.get(page.page_index, []):
                if ann.note_text:
                    pdf.set_font("Helvetica", "I", 10)
                    pdf.multi_cell(0, 5, f"[{ann.type}/{ann.label}] {ann.note_text}")
            pdf.ln(3)

        import tempfile, base64, os
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp_path = tmp.name
        pdf.output(tmp_path)
        with open(tmp_path, "rb") as f:
            data = base64.b64encode(f.read()).decode()
        os.unlink(tmp_path)

        return {"format": "pdf", "filename": f"{doc.title}.pdf", "content_base64": data}
    except ImportError:
        # Fallback: return text
        return _export_txt(doc, pages, annotations)
