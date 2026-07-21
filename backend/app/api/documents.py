"""Document import, text extraction, and management endpoints."""

import json
import os
import uuid
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.config import get_settings
from app.models.database import get_session
from app.models.models import Document, Page

router = APIRouter()

ALLOWED_EXTENSIONS = {"pdf", "docx", "epub", "txt", "png", "jpg", "jpeg"}


def _extract_text(filepath: str, fmt: str) -> tuple[int, list[dict]]:
    """Extract text from a document. Returns (page_count, pages_with_text)."""
    pages = []

    if fmt == "txt":
        with open(filepath, encoding="utf-8", errors="replace") as f:
            text = f.read()
        pages.append({"page_index": 0, "text": text, "dimensions": '{"width":612,"height":792}'})
        return 1, pages

    if fmt == "pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            for i, page in enumerate(doc):
                text = page.get_text() or ""
                rect = page.rect
                dims = json.dumps({"width": rect.width, "height": rect.height})
                pages.append({"page_index": i, "text": text, "dimensions": dims})
            pc = len(doc)
            doc.close()
            return pc, pages
        except Exception:
            pass

    if fmt == "docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            pages.append({"page_index": 0, "text": text, "dimensions": '{"width":612,"height":792}'})
            return 1, pages
        except Exception:
            pass

    if fmt in ("epub",):
        try:
            import ebooklib
            from ebooklib import epub
            book = epub.read_epub(filepath)
            texts = []
            for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(item.get_body_content(), "html.parser")
                texts.append(soup.get_text())
            full = "\n\n".join(texts)
            pages.append({"page_index": 0, "text": full, "dimensions": '{"width":612,"height":792}'})
            return 1, pages
        except Exception:
            pass

    # Fallback: try to read as plain text
    try:
        with open(filepath, encoding="utf-8", errors="replace") as f:
            text = f.read()
        if text.strip():
            pages.append({"page_index": 0, "text": text, "dimensions": '{"width":612,"height":792}'})
            return 1, pages
    except Exception:
        pass

    # Absolute fallback
    pages.append({"page_index": 0, "text": f"[Binary document: {fmt}]", "dimensions": '{"width":612,"height":792}'})
    return 1, pages


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    session: AsyncSession = Depends(get_session),
):
    """Upload a document, extract text, and store in DB."""
    if not file.filename:
        raise HTTPException(400, "No filename provided")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"Unsupported file type: .{ext}")

    settings = get_settings()
    os.makedirs(settings.storage_path, exist_ok=True)

    doc_id = uuid.uuid4().hex
    safe_name = f"{doc_id}_{file.filename}"
    filepath = os.path.join(settings.storage_path, safe_name)

    # Save the uploaded file
    content = await file.read()
    with open(filepath, "wb") as f:
        f.write(content)

    # Extract text
    page_count, pages_data = _extract_text(filepath, ext)

    now = datetime.now(timezone.utc)
    document = Document(
        id=doc_id,
        title=file.filename.rsplit(".", 1)[0] if "." in file.filename else file.filename,
        source_format=ext,
        source_filename=file.filename,
        page_count=page_count,
        storage_path=filepath,
        created_at=now,
        updated_at=now,
    )
    session.add(document)

    for pd in pages_data:
        page = Page(
            document_id=doc_id,
            page_index=pd["page_index"],
            text=pd["text"],
            dimensions=pd["dimensions"],
            ocr_json="{}",
        )
        session.add(page)

    await session.commit()

    return {
        "document_id": doc_id,
        "title": document.title,
        "source_format": ext,
        "page_count": page_count,
        "filename": file.filename,
    }


@router.get("/")
async def list_documents(session: AsyncSession = Depends(get_session)):
    """List all documents with metadata."""
    result = await session.execute(
        select(Document).order_by(Document.created_at.desc())
    )
    docs = result.scalars().all()
    return {
        "documents": [
            {
                "id": d.id,
                "title": d.title,
                "author": d.author,
                "language": d.language,
                "source_format": d.source_format,
                "source_filename": d.source_filename,
                "page_count": d.page_count,
                "created_at": d.created_at.isoformat() if d.created_at else None,
                "updated_at": d.updated_at.isoformat() if d.updated_at else None,
            }
            for d in docs
        ]
    }


@router.get("/{document_id}")
async def get_document(
    document_id: str,
    session: AsyncSession = Depends(get_session),
):
    """Get document details with pages."""
    result = await session.execute(
        select(Document)
        .options(selectinload(Document.pages))
        .where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    return {
        "id": doc.id,
        "title": doc.title,
        "author": doc.author,
        "language": doc.language,
        "source_format": doc.source_format,
        "source_filename": doc.source_filename,
        "page_count": doc.page_count,
        "storage_path": doc.storage_path,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
        "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
        "pages": [
            {
                "id": p.id,
                "page_index": p.page_index,
                "text": p.text,
                "image_path": p.image_path,
                "ocr_json": p.ocr_json,
                "dimensions": p.dimensions,
            }
            for p in (doc.pages or [])
        ],
    }


@router.get("/{document_id}/pages/{page_index}")
async def get_page(
    document_id: str,
    page_index: int,
    session: AsyncSession = Depends(get_session),
):
    """Get a specific page with text data."""
    # Verify document exists
    result = await session.execute(
        select(Document).where(Document.id == document_id)
    )
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(404, "Document not found")

    # Get specific page
    result = await session.execute(
        select(Page).where(
            Page.document_id == document_id,
            Page.page_index == page_index,
        )
    )
    page = result.scalar_one_or_none()
    if not page:
        raise HTTPException(404, f"Page {page_index} not found")

    return {
        "id": page.id,
        "document_id": page.document_id,
        "page_index": page.page_index,
        "text": page.text,
        "image_path": page.image_path,
        "ocr_json": page.ocr_json,
        "dimensions": page.dimensions,
    }
